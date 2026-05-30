from fastapi import (
    FastAPI,
    UploadFile,
    File,
    Form,
    HTTPException,
    Depends,
)

from fastapi.responses import (
    FileResponse,
    JSONResponse,
)

from fastapi.middleware.cors import CORSMiddleware

from sqlalchemy.orm import Session

from dotenv import load_dotenv

from openai import OpenAI

from langchain_community.document_loaders import PyPDFLoader

from docx import Document

from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
)

from reportlab.lib.styles import getSampleStyleSheet

from openpyxl import Workbook

from pydantic import BaseModel

from datetime import datetime

from database import (
    SessionLocal,
    engine,
    Base,
)

from models import (
    User,
    Article,
    Review,
)

from auth import (
    hash_password,
    verify_password,
    create_access_token,
    get_current_user,
    admin_required,
)

import traceback
import shutil
import os
import re


load_dotenv()


app = FastAPI(
    title="AI Academic Reviewer",
    version="6.0",
)


app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


Base.metadata.create_all(bind=engine)


OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")

client = OpenAI(
    api_key=OPENAI_API_KEY,
)


uploaded_pdf_text = ""
uploaded_filename = ""
last_article_review = ""


class ChatRequest(BaseModel):
    prompt: str | None = None
    message: str | None = None


class CompareRequest(BaseModel):
    original_text: str
    corrected_text: str


def get_db():
    db = SessionLocal()

    try:
        yield db

    finally:
        db.close()


def blind_review_text(text: str):
    patterns = [
        r"(?i)autor:",
        r"(?i)authors:",
        r"(?i)correo:",
        r"(?i)e-mail:",
        r"(?i)universidad",
        r"(?i)instituto",
    ]

    cleaned = text

    for pattern in patterns:
        cleaned = re.sub(
            pattern,
            "[BLIND]",
            cleaned,
        )

    return cleaned


def calculate_score(review_text: str):
    review_lower = review_text.lower()

    if "rechazado" in review_lower:
        return 40

    if "cambios mayores" in review_lower:
        return 55

    if "cambios menores" in review_lower:
        return 75

    if "aceptado" in review_lower:
        return 90

    return 60


def detect_badge(review_text: str):
    review_lower = review_text.lower()

    if "rechazado" in review_lower:
        return "Rechazado"

    if "cambios mayores" in review_lower:
        return "Requiere cambios mayores"

    if "cambios menores" in review_lower:
        return "Aceptado con cambios menores"

    if "aceptado" in review_lower:
        return "Aceptado sin cambios"

    return "Requiere cambios mayores"


def detect_ai_probability(review_text: str):
    review_lower = review_text.lower()

    if "muy artificial" in review_lower:
        return "Alta"

    if "moderadamente artificial" in review_lower:
        return "Media"

    return "Baja"


def status_from_badge(badge: str):
    mapping = {
        "Aceptado sin cambios": "accepted",
        "Aceptado con cambios menores": "minor_revision",
        "Requiere cambios mayores": "major_revision",
        "Rechazado": "rejected",
    }

    return mapping.get(
        badge,
        "under_review",
    )


@app.get("/")
def root():
    return {
        "message": "AI Academic Reviewer API funcionando",
        "status": "online",
    }
@app.post("/register")
def register(
    username: str = Form(...),
    password: str = Form(...),
    role: str = Form("reviewer"),
    db: Session = Depends(get_db),
):
    try:
        existing_user = (
            db.query(User)
            .filter(User.username == username)
            .first()
        )

        if existing_user:
            raise HTTPException(
                status_code=400,
                detail="El usuario ya existe",
            )

        hashed_password = hash_password(password)

        new_user = User(
            username=username,
            password=hashed_password,
            role=role,
        )

        db.add(new_user)
        db.commit()
        db.refresh(new_user)

        return {
            "message": "Usuario registrado correctamente",
            "username": new_user.username,
            "role": new_user.role,
        }

    except HTTPException:
        raise

    except Exception as e:
        traceback.print_exc()

        return JSONResponse(
            status_code=500,
            content={
                "error": str(e)
            },
        )


@app.post("/login")
def login(
    username: str = Form(...),
    password: str = Form(...),
    db: Session = Depends(get_db),
):
    try:
        user = (
            db.query(User)
            .filter(User.username == username)
            .first()
        )

        if not user:
            raise HTTPException(
                status_code=401,
                detail="Usuario incorrecto",
            )

        if not verify_password(
            password,
            user.password,
        ):
            raise HTTPException(
                status_code=401,
                detail="Contraseña incorrecta",
            )

        token = create_access_token(
            {
                "sub": user.username,
                "role": user.role,
            }
        )

        return {
            "access_token": token,
            "token_type": "bearer",
            "username": user.username,
            "role": user.role,
        }

    except HTTPException:
        raise

    except Exception as e:
        traceback.print_exc()

        return JSONResponse(
            status_code=500,
            content={
                "error": str(e)
            },
        )


@app.post("/upload-pdf")
def upload_pdf(
    file: UploadFile = File(...),
):
    global uploaded_pdf_text
    global uploaded_filename
    global last_article_review

    try:
        os.makedirs(
            "uploaded_files",
            exist_ok=True,
        )

        file_path = f"uploaded_files/{file.filename}"

        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(
                file.file,
                buffer,
            )

        loader = PyPDFLoader(file_path)
        documents = loader.load()

        uploaded_pdf_text = "\n\n".join(
            [
                document.page_content
                for document in documents
            ]
        )

        uploaded_filename = file.filename
        last_article_review = ""

        return {
            "message": "PDF cargado correctamente",
            "filename": file.filename,
            "pages": len(documents),
        }

    except Exception as e:
        traceback.print_exc()

        return JSONResponse(
            status_code=500,
            content={
                "error": str(e)
            },
        )


@app.post("/ask-pdf")
def ask_pdf(
    question: str = Form(...),
):
    global uploaded_pdf_text

    try:
        if not uploaded_pdf_text:
            raise HTTPException(
                status_code=400,
                detail="Primero debes subir un PDF",
            )

        context = uploaded_pdf_text[:25000]

        prompt = f"""
Responde únicamente con base en el siguiente documento.

DOCUMENTO:
{context}

PREGUNTA:
{question}
"""

        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {
                    "role": "system",
                    "content": (
                        "Eres un asistente académico experto en análisis documental."
                    ),
                },
                {
                    "role": "user",
                    "content": prompt,
                },
            ],
        )

        answer = response.choices[0].message.content

        return {
            "answer": answer,
            "sources": [
                {
                    "page": 1,
                    "content": uploaded_pdf_text[:300],
                }
            ],
        }

    except HTTPException:
        raise

    except Exception as e:
        traceback.print_exc()

        return JSONResponse(
            status_code=500,
            content={
                "error": str(e)
            },
        )
@app.post("/articles")
def create_article(
    title: str = Form(...),
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    try:
        os.makedirs(
            "uploaded_files",
            exist_ok=True,
        )

        file_path = f"uploaded_files/{file.filename}"

        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(
                file.file,
                buffer,
            )

        article = Article(
            title=title,
            filename=file.filename,
            status="submitted",
        )

        db.add(article)
        db.commit()
        db.refresh(article)

        return {
            "message": "Artículo creado correctamente",
            "article_id": article.id,
            "title": article.title,
            "filename": article.filename,
            "status": article.status,
        }

    except Exception as e:
        traceback.print_exc()

        return JSONResponse(
            status_code=500,
            content={
                "error": str(e)
            },
        )


@app.get("/articles")
def get_articles(
    db: Session = Depends(get_db),
):
    articles = (
        db.query(Article)
        .order_by(Article.created_at.desc())
        .all()
    )

    return [
        {
            "id": article.id,
            "title": article.title,
            "filename": article.filename,
            "status": article.status,
            "created_at": article.created_at,
            "reviews_count": len(article.reviews),
        }
        for article in articles
    ]


@app.get("/articles/{article_id}")
def get_article(
    article_id: int,
    db: Session = Depends(get_db),
):
    article = (
        db.query(Article)
        .filter(Article.id == article_id)
        .first()
    )

    if not article:
        raise HTTPException(
            status_code=404,
            detail="Artículo no encontrado",
        )

    return {
        "id": article.id,
        "title": article.title,
        "filename": article.filename,
        "status": article.status,
        "created_at": article.created_at,
        "reviews": [
            {
                "id": review.id,
                "reviewer_id": review.reviewer_id,
                "review_type": review.review_type,
                "score": review.score,
                "badge": review.badge,
                "ai_probability": review.ai_probability,
                "created_at": review.created_at,
            }
            for review in article.reviews
        ],
    }


@app.post("/articles/{article_id}/assign")
def assign_reviewer(
    article_id: int,
    reviewer_id: int = Form(...),
    current_user: User = Depends(admin_required),
    db: Session = Depends(get_db),
):
    article = (
        db.query(Article)
        .filter(Article.id == article_id)
        .first()
    )

    reviewer = (
        db.query(User)
        .filter(User.id == reviewer_id)
        .first()
    )

    if not article:
        raise HTTPException(
            status_code=404,
            detail="Artículo no encontrado",
        )

    if not reviewer:
        raise HTTPException(
            status_code=404,
            detail="Revisor no encontrado",
        )

    article.status = "under_review"
    db.commit()

    return {
        "message": "Revisor asignado",
        "article": article.title,
        "reviewer": reviewer.username,
        "status": article.status,
    }


@app.post("/articles/{article_id}/status")
def update_article_status(
    article_id: int,
    status: str = Form(...),
    current_user: User = Depends(admin_required),
    db: Session = Depends(get_db),
):
    allowed_statuses = [
        "submitted",
        "under_review",
        "minor_revision",
        "major_revision",
        "accepted",
        "rejected",
        "published",
    ]

    if status not in allowed_statuses:
        raise HTTPException(
            status_code=400,
            detail="Estado editorial no válido",
        )

    article = (
        db.query(Article)
        .filter(Article.id == article_id)
        .first()
    )

    if not article:
        raise HTTPException(
            status_code=404,
            detail="Artículo no encontrado",
        )

    article.status = status
    db.commit()

    return {
        "message": "Estado actualizado",
        "article_id": article.id,
        "status": article.status,
    }


@app.post("/review-article")
def review_article(
    review_type: str = Form(...),
    blind_review: bool = Form(...),
    article_id: int | None = Form(None),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    global uploaded_pdf_text
    global uploaded_filename
    global last_article_review

    try:
        if not uploaded_pdf_text:
            raise HTTPException(
                status_code=400,
                detail="Primero debes subir un artículo PDF",
            )

        article_text = uploaded_pdf_text[:30000]

        if blind_review:
            article_text = blind_review_text(article_text)

        prompt = f"""
Eres un sistema multiagente de arbitraje académico especializado en evaluación científica.

Tipo de evaluación:
{review_type}

Revisión ciega:
{blind_review}

Actúa como:
1. Revisor metodológico
2. Revisor teórico
3. Revisor editorial
4. Revisor APA
5. Editor en jefe

ARTÍCULO:
{article_text}

Genera un dictamen PROFUNDO, CRÍTICO, CONSTRUCTIVO y ACADÉMICO.

Usa exactamente esta estructura:

# Dictamen académico multiagente

# Badge de dictamen editorial

Elige solo una opción:
- Aceptado sin cambios
- Aceptado con cambios menores
- Requiere cambios mayores
- Rechazado

# Score general del artículo

# Revisión metodológica

# Revisión teórica

# Revisión editorial y de redacción

# Revisión APA y formato

# Tabla sintética de observaciones

Usa tabla Markdown:
| Apartado | Problema | Recomendación | Prioridad |
|---|---|---|---|

# Fortalezas

# Debilidades

# Recomendaciones concretas

# Dictamen final del editor en jefe
"""

        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {
                    "role": "system",
                    "content": (
                        "Eres un árbitro científico riguroso, humano y constructivo."
                    ),
                },
                {
                    "role": "user",
                    "content": prompt,
                },
            ],
        )

        review_text = response.choices[0].message.content

        score = calculate_score(review_text)
        badge = detect_badge(review_text)
        ai_probability = detect_ai_probability(review_text)

        article = None

        if article_id:
            article = (
                db.query(Article)
                .filter(Article.id == article_id)
                .first()
            )

        if not article:
            article = Article(
                title=uploaded_filename or "Artículo sin título",
                filename=uploaded_filename or "uploaded_article.pdf",
                status="under_review",
            )

            db.add(article)
            db.commit()
            db.refresh(article)

        new_review = Review(
            article_id=article.id,
            reviewer_id=current_user.id,
            filename=uploaded_filename or "uploaded_article.pdf",
            review_type=review_type,
            review_content=review_text,
            score=int(score),
            ai_probability=ai_probability,
            badge=badge,
            created_at=datetime.utcnow(),
        )

        article.status = status_from_badge(badge)

        db.add(new_review)
        db.commit()
        db.refresh(new_review)

        last_article_review = review_text

        return {
            "review": review_text,
            "review_content": review_text,
            "dictamen": review_text,
            "score": str(score),
            "badge": badge,
            "ai_probability": ai_probability,
            "review_id": new_review.id,
            "article_id": article.id,
            "article_status": article.status,
        }

    except HTTPException:
        raise

    except Exception as e:
        traceback.print_exc()

        return JSONResponse(
            status_code=500,
            content={
                "error": str(e)
            },
        )
@app.post("/chat")
def chat(
    data: ChatRequest,
):
    user_message = data.prompt or data.message or ""

    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {
                "role": "system",
                "content": "Eres un asistente académico claro y profesional.",
            },
            {
                "role": "user",
                "content": user_message,
            },
        ],
    )

    return {
        "response": response.choices[0].message.content
    }


@app.post("/compare")
def compare_versions_json(
    request: CompareRequest,
):
    return compare_logic(
        request.original_text,
        request.corrected_text,
    )


@app.post("/compare-reviews")
def compare_versions_form(
    original_text: str = Form(...),
    corrected_text: str = Form(...),
):
    return compare_logic(
        original_text,
        corrected_text,
    )


def compare_logic(
    original_text: str,
    corrected_text: str,
):
    prompt = f"""
Compara ambas versiones de un artículo académico.

VERSIÓN ORIGINAL:
{original_text}

VERSIÓN CORREGIDA:
{corrected_text}

Evalúa:
1. Qué observaciones fueron atendidas.
2. Qué problemas persisten.
3. Qué mejoró.
4. Qué sigue faltando.
5. Nivel de mejora general.

Redacta en formato académico.
"""

    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {
                "role": "system",
                "content": (
                    "Eres un experto en evaluación académica y mejora de manuscritos."
                ),
            },
            {
                "role": "user",
                "content": prompt,
            },
        ],
    )

    return {
        "comparison": response.choices[0].message.content
    }


@app.get("/reviews")
def get_reviews(
    db: Session = Depends(get_db),
):
    reviews = (
        db.query(Review)
        .order_by(Review.created_at.desc())
        .all()
    )

    return [
        {
            "id": review.id,
            "article_id": review.article_id,
            "reviewer_id": review.reviewer_id,
            "filename": review.filename,
            "review_type": review.review_type,
            "badge": review.badge,
            "score": str(review.score),
            "ai_probability": review.ai_probability,
            "created_at": review.created_at,
        }
        for review in reviews
    ]


@app.get("/review/{review_id}")
def get_review_alias(
    review_id: int,
    db: Session = Depends(get_db),
):
    return get_review(
        review_id,
        db,
    )


@app.get("/reviews/{review_id}")
def get_review(
    review_id: int,
    db: Session = Depends(get_db),
):
    review = (
        db.query(Review)
        .filter(Review.id == review_id)
        .first()
    )

    if not review:
        raise HTTPException(
            status_code=404,
            detail="Revisión no encontrada",
        )

    return {
        "id": review.id,
        "article_id": review.article_id,
        "reviewer_id": review.reviewer_id,
        "filename": review.filename,
        "review_type": review.review_type,
        "badge": review.badge,
        "score": str(review.score),
        "review": review.review_content,
        "review_content": review.review_content,
        "dictamen": review.review_content,
        "ai_probability": review.ai_probability,
        "created_at": review.created_at,
    }


@app.delete("/reviews/{review_id}")
def delete_review(
    review_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    review = (
        db.query(Review)
        .filter(Review.id == review_id)
        .first()
    )

    if not review:
        raise HTTPException(
            status_code=404,
            detail="Dictamen no encontrado",
        )

    db.delete(review)
    db.commit()

    return {
        "message": "Dictamen eliminado correctamente",
        "review_id": review_id,
    }


@app.get("/dashboard")
def dashboard(
    db: Session = Depends(get_db),
):
    reviews = db.query(Review).all()
    articles = db.query(Article).all()

    return {
        "total_reviews": len(reviews),
        "total_articles": len(articles),
        "accepted": len(
            [
                review
                for review in reviews
                if review.badge == "Aceptado sin cambios"
            ]
        ),
        "minor_changes": len(
            [
                review
                for review in reviews
                if review.badge == "Aceptado con cambios menores"
            ]
        ),
        "major_changes": len(
            [
                review
                for review in reviews
                if review.badge == "Requiere cambios mayores"
            ]
        ),
        "rejected": len(
            [
                review
                for review in reviews
                if review.badge == "Rechazado"
            ]
        ),
        "submitted_articles": len(
            [
                article
                for article in articles
                if article.status == "submitted"
            ]
        ),
        "under_review_articles": len(
            [
                article
                for article in articles
                if article.status == "under_review"
            ]
        ),
        "published_articles": len(
            [
                article
                for article in articles
                if article.status == "published"
            ]
        ),
    }


@app.get("/users")
def get_users(
    current_user: User = Depends(admin_required),
    db: Session = Depends(get_db),
):
    users = db.query(User).all()

    return [
        {
            "id": user.id,
            "username": user.username,
            "role": user.role,
        }
        for user in users
    ]


@app.post("/users/{user_id}/role")
def update_user_role(
    user_id: int,
    role: str = Form(...),
    current_user: User = Depends(admin_required),
    db: Session = Depends(get_db),
):
    allowed_roles = [
        "admin",
        "editor",
        "reviewer",
        "author",
    ]

    if role not in allowed_roles:
        raise HTTPException(
            status_code=400,
            detail="Rol no válido",
        )

    user = (
        db.query(User)
        .filter(User.id == user_id)
        .first()
    )

    if not user:
        raise HTTPException(
            status_code=404,
            detail="Usuario no encontrado",
        )

    user.role = role
    db.commit()

    return {
        "message": "Rol actualizado",
        "username": user.username,
        "role": user.role,
    }
def build_word_file(
    review_text: str,
    file_path: str = "dictamen_academico.docx",
):
    document = Document()

    document.add_heading(
        "Dictamen académico",
        level=1,
    )

    document.add_paragraph(
        f"Fecha: {datetime.now().strftime('%d/%m/%Y %H:%M')}"
    )

    for line in review_text.split("\n"):
        clean = line.strip()

        if not clean:
            continue

        if clean.startswith("# "):
            document.add_heading(
                clean.replace("# ", ""),
                level=1,
            )

        elif clean.startswith("## "):
            document.add_heading(
                clean.replace("## ", ""),
                level=2,
            )

        else:
            document.add_paragraph(clean)

    document.save(file_path)

    return file_path


def build_pdf_file(
    review_text: str,
    file_path: str = "dictamen_academico.pdf",
):
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.units import cm
    from reportlab.platypus import Table, TableStyle
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
    from reportlab.lib.styles import ParagraphStyle

    doc = SimpleDocTemplate(
        file_path,
        pagesize=letter,
        rightMargin=2 * cm,
        leftMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
    )

    styles = getSampleStyleSheet()
    elements = []

    title_style = ParagraphStyle(
        "InstitutionalTitle",
        parent=styles["Heading1"],
        alignment=TA_CENTER,
        fontName="Helvetica-Bold",
        fontSize=16,
        textColor=colors.HexColor("#0B1B33"),
        spaceAfter=10,
    )

    subtitle_style = ParagraphStyle(
        "InstitutionalSubtitle",
        parent=styles["Normal"],
        alignment=TA_CENTER,
        fontName="Helvetica",
        fontSize=10,
        textColor=colors.HexColor("#4B5563"),
        spaceAfter=16,
    )

    section_style = ParagraphStyle(
        "SectionTitle",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=12,
        textColor=colors.HexColor("#0B1B33"),
        spaceBefore=12,
        spaceAfter=6,
    )

    body_style = ParagraphStyle(
        "BodyJustified",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=9.5,
        leading=13,
        alignment=TA_JUSTIFY,
        textColor=colors.HexColor("#111827"),
        spaceAfter=6,
    )

    small_style = ParagraphStyle(
        "Small",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=8,
        leading=10,
        textColor=colors.HexColor("#374151"),
    )

    header_table = Table(
        [
            [
                Paragraph("<b>Universidad Veracruzana</b>", title_style),
            ],
            [
                Paragraph(
                    "Instituto de Investigaciones en Contaduría",
                    subtitle_style,
                ),
            ],
            [
                Paragraph(
                    "Dictamen académico asistido por inteligencia artificial",
                    subtitle_style,
                ),
            ],
        ],
        colWidths=[17 * cm],
    )

    header_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F3F6FA")),
                ("BOX", (0, 0), (-1, -1), 1, colors.HexColor("#D1D5DB")),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 8),
            ]
        )
    )

    elements.append(header_table)
    elements.append(Spacer(1, 14))

    info_data = [
        ["Fecha de generación", datetime.now().strftime("%d/%m/%Y %H:%M")],
        ["Tipo de documento", "Dictamen académico"],
        ["Sistema", "AI Academic Reviewer"],
        ["Institución", "Universidad Veracruzana"],
    ]

    info_table = Table(
        info_data,
        colWidths=[5 * cm, 12 * cm],
    )

    info_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#0B1B33")),
                ("TEXTCOLOR", (0, 0), (0, -1), colors.white),
                ("BACKGROUND", (1, 0), (1, -1), colors.HexColor("#FFFFFF")),
                ("TEXTCOLOR", (1, 0), (1, -1), colors.HexColor("#111827")),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#D1D5DB")),
                ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
                ("FONTSIZE", (0, 0), (-1, -1), 8.5),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )

    elements.append(info_table)
    elements.append(Spacer(1, 16))

    elements.append(
        Paragraph(
            "Contenido del dictamen",
            section_style,
        )
    )

    for line in review_text.split("\n"):
        clean = line.strip()

        if not clean:
            elements.append(Spacer(1, 5))
            continue

        safe_line = (
            clean.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
        )

        if clean.startswith("# "):
            elements.append(
                Paragraph(
                    safe_line.replace("# ", ""),
                    section_style,
                )
            )

        elif clean.startswith("## "):
            elements.append(
                Paragraph(
                    safe_line.replace("## ", ""),
                    section_style,
                )
            )

        elif clean.startswith("|"):
            elements.append(
                Paragraph(
                    safe_line,
                    small_style,
                )
            )

        else:
            elements.append(
                Paragraph(
                    safe_line,
                    body_style,
                )
            )

    elements.append(Spacer(1, 20))

    footer = Paragraph(
        "Nota: Este dictamen fue generado como apoyo académico mediante herramientas de inteligencia artificial. "
        "La decisión editorial final debe ser validada por el comité académico correspondiente.",
        small_style,
    )

    elements.append(footer)

    doc.build(elements)

    return file_path


@app.post("/export-review")
def export_review():
    return export_review_word()


@app.post("/export-review-word")
def export_review_word():
    global last_article_review

    if not last_article_review:
        raise HTTPException(
            status_code=400,
            detail="Primero debes generar un dictamen",
        )

    file_path = build_word_file(last_article_review)

    return FileResponse(
        path=file_path,
        filename="dictamen_academico.docx",
        media_type=(
            "application/vnd.openxmlformats-officedocument."
            "wordprocessingml.document"
        ),
    )


@app.get("/export-review-word")
def export_review_word_get():
    return export_review_word()


@app.post("/export-review-pdf")
def export_review_pdf():
    global last_article_review

    if not last_article_review:
        raise HTTPException(
            status_code=400,
            detail="Primero debes generar un dictamen",
        )

    file_path = build_pdf_file(last_article_review)

    return FileResponse(
        path=file_path,
        filename="dictamen_academico.pdf",
        media_type="application/pdf",
    )


@app.get("/export-review-pdf")
def export_review_pdf_get():
    return export_review_pdf()


@app.get("/reviews/{review_id}/pdf")
def download_review_pdf(
    review_id: int,
    db: Session = Depends(get_db),
):
    review = (
        db.query(Review)
        .filter(Review.id == review_id)
        .first()
    )

    if not review:
        raise HTTPException(
            status_code=404,
            detail="Dictamen no encontrado",
        )

    file_path = build_pdf_file(
        review.review_content,
        file_path=f"dictamen_{review.id}.pdf",
    )

    return FileResponse(
        path=file_path,
        filename=f"dictamen_{review.id}.pdf",
        media_type="application/pdf",
    )


@app.get("/reviews/{review_id}/word")
def download_review_word(
    review_id: int,
    db: Session = Depends(get_db),
):
    review = (
        db.query(Review)
        .filter(Review.id == review_id)
        .first()
    )

    if not review:
        raise HTTPException(
            status_code=404,
            detail="Dictamen no encontrado",
        )

    file_path = build_word_file(
        review.review_content,
        file_path=f"dictamen_{review.id}.docx",
    )

    return FileResponse(
        path=file_path,
        filename=f"dictamen_{review.id}.docx",
        media_type=(
            "application/vnd.openxmlformats-officedocument."
            "wordprocessingml.document"
        ),
    )


@app.get("/reviews-export/excel")
def export_reviews_excel(
    db: Session = Depends(get_db),
):
    reviews = (
        db.query(Review)
        .order_by(Review.created_at.desc())
        .all()
    )

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Historial de dictámenes"

    headers = [
        "ID",
        "Artículo ID",
        "Revisor ID",
        "Archivo",
        "Tipo de revisión",
        "Score",
        "Dictamen",
        "IA",
        "Fecha",
    ]

    sheet.append(headers)

    for review in reviews:
        sheet.append(
            [
                review.id,
                review.article_id,
                review.reviewer_id,
                review.filename,
                review.review_type,
                review.score,
                review.badge,
                review.ai_probability,
                review.created_at.strftime("%d/%m/%Y %H:%M")
                if review.created_at
                else "",
            ]
        )

    file_path = "historial_dictamenes.xlsx"
    workbook.save(file_path)

    return FileResponse(
        path=file_path,
        filename="historial_dictamenes.xlsx",
        media_type=(
            "application/vnd.openxmlformats-officedocument."
            "spreadsheetml.sheet"
        ),
    )


@app.get("/dashboard-export/pdf")
def export_dashboard_pdf(
    db: Session = Depends(get_db),
):
    reviews = db.query(Review).all()
    articles = db.query(Article).all()

    content = f"""
# Reporte general del sistema

Total de artículos: {len(articles)}
Total de dictámenes: {len(reviews)}
Aceptados: {len([r for r in reviews if r.badge == "Aceptado sin cambios"])}
Cambios menores: {len([r for r in reviews if r.badge == "Aceptado con cambios menores"])}
Cambios mayores: {len([r for r in reviews if r.badge == "Requiere cambios mayores"])}
Rechazados: {len([r for r in reviews if r.badge == "Rechazado"])}

# Estadísticas de IA

Riesgo IA alto: {len([r for r in reviews if r.ai_probability == "Alta"])}
Riesgo IA medio: {len([r for r in reviews if r.ai_probability == "Media"])}
Riesgo IA bajo: {len([r for r in reviews if r.ai_probability == "Baja"])}
"""

    file_path = build_pdf_file(
        content,
        file_path="reporte_dashboard.pdf",
    )

    return FileResponse(
        path=file_path,
        filename="reporte_dashboard.pdf",
        media_type="application/pdf",
    )